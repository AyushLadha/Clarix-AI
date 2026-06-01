import os
import io
import json
import pandas as pd
from dotenv import load_dotenv
from pyparsing import col
import streamlit as st
import matplotlib
import matplotlib.pyplot as plt
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_core.callbacks import BaseCallbackHandler
from langchain.agents import create_agent
from langchain_core.tools import tool
from langgraph.checkpoint.memory import MemorySaver
import uuid
import chromadb
import tempfile
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader 

# --------------- Load API key ---------------
load_dotenv()
# Works both locally (.env) and on Streamlit Cloud (secrets)
api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
    except:
        st.error("GEMINI_API_KEY not found. Please set it in your .env file or Streamlit secrets.")
        st.stop()

# --------------- Page Configuration ---------------
st.set_page_config(
    page_title = "Clarix", 
    page_icon = "✦", 
    layout = "wide"
    )

# --------------- Agent Callback Handler ---------------
class StreamlitAgentCallback(BaseCallbackHandler):
    """Show live agent steps in streamlit status container"""

    def __init__(self, status_container):
        self.status = status_container

    def on_tool_start(self, serialized, input_str, **kwargs):
        try:
            tool_name = serialized.get('name', 'unknown')
            if tool_name == "get_dataframe_info":
                self.status.update(label = "📋 Reading dataset structure...")
            elif tool_name == "query_dataframe":
                self.status.update(label = "⚡ Running data query...")
            elif tool_name == "get_column_values":
                self.status.update(label = "🔎 Checking column values...")
            elif tool_name == "generate_chart":
                self.status.update(label = "📈 Generating chart...")
        except:
            pass # silently ignore threading context errors

    def on_tool_end(self, output, **kwargs):
        try:
            self.status.update(label = "🔄 Processing results...")
        except:
            pass

    def on_llm_start(self, serialized, prompts, **kwargs):
        try:
            self.status.update(label="🤔 Thinking...")
        except:
            pass

    def on_agent_finish(self, finish, **kwargs):
        try:
            self.status.update(label="✅ Done!", state="complete", expanded=False)
        except:
            pass

# --------------- Helper function to clean dataframe columns ---------------
def clean_dataframe(df):
    """Fix mixed type columns that cause Arrow/PyArrow errors in Streamlit"""
    for col in df.columns:
        # If column is object type but has mixed int/str values — convert all to str
        if df[col].dtype == 'object':
            df[col] = df[col].astype(str)
    return df

# --------------- Function to summarize the data ---------------
def summarize_dataframe(df):
    lines = []
    # Basic shape, column names and numeric stats, categorical columns and missing values
    lines.append(f"Dataset: {df.shape[0]:,} rows x {df.shape[1]} columns")
                
    lines.append(f"\nColumns: {df.columns.tolist()}")
                
    numeric_columns = df.select_dtypes(include = 'number')
    if not numeric_columns.empty:
        lines.append(f"\nNumeric Statistics: {numeric_columns.describe().round(2).to_string() }")
                
    categorical_columns = df.select_dtypes(include = ['object', 'string']).columns.tolist()

     # For categorical columns, we want to show the distribution of values for the top categorical columns, 
     # so that we can get a sense of what the data looks like and what kind of insights we might be able to generate from it.
    for col in categorical_columns:
        top = df[col].value_counts().head(10).to_dict()
        lines.append(f"\nTransaction/record count by '{col}':\n{top}")

    # For grouping and aggregation, we want to pick categorical columns that have a reasonable number of unique values (not too many, not too few) 
    # to avoid overwhelming the model with too much information or having meaningless groupings.
    good_group_cols = [
        col for col in categorical_columns
        if 2 <= df[col].nunique() <= 30  # sweet spot for grouping
    ]
    # For aggregation columns, we want to pick numeric columns that have a good amount of variability and are not just identifiers or constant values, 
    # as those would not provide meaningful insights when aggregated.
    good_agg_cols = numeric_columns.columns.tolist()

    # We will show some sample aggregations by key groups to give the model a sense of what kind of insights can be generated from the data, 
    # and to provide it with some specific examples of trends and patterns in the data that it can reference when generating the report.
    if good_group_cols and good_agg_cols:
        lines.append(f"\n--- Aggregations by Key Groups ---")
        for group_col in good_group_cols[:5]:   # top 5 grouping cols
            for agg_col in good_agg_cols[:5]:   # top 5 numeric cols
                try:
                    agg = df.groupby(group_col)[agg_col].agg(['sum', 'mean', 'count'])
                    agg = agg.sort_values('sum', ascending=False).head(10).round(2)
                    lines.append(f"\n{agg_col} by {group_col}:\n{agg.to_string()}")
                except:
                    pass
                
    nulls = df.isnull().sum()
    if nulls.any():
        lines.append(f"\nMissing values: {nulls[nulls > 0].to_dict()}")
                
    lines.append(f"\nSample rows: {df.head(5).to_string()}")
                
    return "\n".join(lines)

# --------------- Function for generating charts using Gemini ---------------
def get_chart_columns(df, llm):
    # We will use the LLM to help us identify which columns might be most interesting to generate charts for, based on the data summary and the column names. 
    # This way, we can leverage the model's understanding of the data and its ability to identify patterns and trends to guide our chart generation, 
    # rather than just relying on arbitrary rules for picking columns.
    """Ask the LLM to identify which columns in the dataset would be most interesting and insightful to generate charts for, based on the data summary and column names."""

    numeric_columns = df.select_dtypes(include = 'number').columns.tolist()
    categorical_columns = df.select_dtypes(include = ['object', 'string']).columns.tolist()

    # Try ot convert object colum to numeric if possible
    for col in categorical_columns:
        try:
            converted = pd.to_numeric(df[col], errors='coerce')
            if converted.notna().sum() > len(df) * 0.5:  # more than 50% convertible
                df[col] = converted
                numeric_columns.append(col)
                categorical_columns.remove(col)
        except:
            pass

    # If still no numeric or categorical columns, skip charting
    if not numeric_columns and not categorical_columns:
        st.info("No suitable columns found for chart generation.")
        return None

    # Handle empty numeric columns in describe
    if numeric_columns:
        numeric_stats = df[numeric_columns].describe().round(2).to_string()
    else:
        numeric_stats = "No numeric columns available"

    col_info = f"""
                Numeric columns: {numeric_columns}
                Categorical columns: {categorical_columns}
                Sample data:
                {df.head(10).to_string()}
                Numeric stats:
                {df[numeric_columns].describe().round(2).to_string() if numeric_columns else "No numeric columns available"}
                """
    messages = [
        SystemMessage(content = """You are a data visualization expert. You are given a summary of a dataset, including the column names, data types, and some sample data.
                      Given a list of columns, you should identify which ones would be most interesting and insightful to generate charts for, based on the data summary and column names.
                      RULES:
                        - NEVER pick ID columns (customer_id, product_key, order_id etc)
                        - NEVER pick code columns (zip_code, postal_code etc)
                        - NEVER pick index columns
                        - NEVER pick date/time columns for scatter or histogram
                        - DO pick metric columns (revenue, profit, salary, sales, quantity, score etc)
                        - DO pick meaningful categorical columns (country, department, category, gender etc)
                        - DO pick time columns ONLY for trend charts (year, month)

                        Respond ONLY in this exact JSON format, no other text:
                        {
                            "histogram_cols": ["col1", "col2", "col3", "col4"],
                            "bar_chart_col": "col1",
                            "scatter_x": "col1",
                            "scatter_y": "col2",
                            "reasoning": "brief explanation"
                        }
                        If you cannot find suitable columns for a particular chart type, return an empty list or null for that field, 
                      but still provide reasoning in the response about why you couldn't find suitable columns for that chart type."""),

        HumanMessage(content = f"Pick the best columns for charts from this dataset:\n\n{col_info}.")
    ]

    response = llm.invoke(messages)
    try:
        import json
        text = response.content[0]['text']
        text = text.replace('```json', '').replace('```', '').strip()
        result = json.loads(text)
        return result
    except Exception as e:
        st.warning("AI could not select chart columns — using fallback method.")
        return None

# --------------- Function to generate charts for a given dataframe and display in Streamlit ---------------
def generate_chart_bytes(df, llm): 
    """
    Generates charts and returns them as a list of (title, bytes) tuples.
    Used by both Streamlit display and Word document generation.
    """

    matplotlib.use('Agg')
    sheet_charts = []

    # Skip if no meaningful column names
    unnamed_cols = [col for col in df.columns
                    if str(col).startswith('Unnamed')
                    or isinstance(col, int)
                    or str(col).replace('.','').replace('-','').isdigit()]
    if len(unnamed_cols) == len(df.columns):
        return []  # return empty — caller handles the message

    chart_cols = get_chart_columns(df, llm)
    if not chart_cols:
        return []

    # Chart 1: Histograms
    hist_cols = [c for c in chart_cols.get('histogram_cols', []) if c in df.columns]
    if hist_cols:
        cols_to_plot = hist_cols[:4]
        fig, axes = plt.subplots(1, len(cols_to_plot), figsize=(4 * len(cols_to_plot), 3))
        if len(cols_to_plot) == 1:
            axes = [axes]
        for ax, col in zip(axes, cols_to_plot):
            df[col].dropna().hist(ax = ax, bins = 20, color = 'skyblue', edgecolor = 'black')
            ax.set_title(col, fontsize = 10)
            ax.set_ylabel("Count", fontsize = 8)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format = 'png', dpi = 150, bbox_inches = 'tight')
        buf.seek(0)
        sheet_charts.append(("Distribution of Numeric Columns", buf.getvalue()))
        plt.close()

    # Charts 2, 3, and 4 side by side
    fig, axes = plt.subplots(1, 3, figsize=(12, 3))

    # Chart 2: Bar chart
    bar_col = chart_cols.get('bar_chart_col')
    if bar_col and bar_col in df.columns:
        df[bar_col].value_counts().head(10).plot(
            kind = 'bar', ax = axes[0], color = 'salmon', edgecolor = 'black'
        )
        axes[0].set_title(f"Top Values in '{bar_col}'", fontsize = 8)
        axes[0].set_ylabel("Count", fontsize = 7)
        axes[0].tick_params(axis = 'x', rotation = 45, labelsize = 7)
    else:
        axes[0].axis('off')
        axes[0].set_title("Bar Chart Not Available", fontsize = 8)

    # Chart 3: Scatter
    scatter_x = chart_cols.get('scatter_x')
    scatter_y = chart_cols.get('scatter_y')
    if scatter_x and scatter_y and scatter_x in df.columns and scatter_y in df.columns:
        scatter_df = df[[scatter_x, scatter_y]].dropna()
        if len(scatter_df) > 0:
            axes[1].scatter(
                scatter_df[scatter_x], scatter_df[scatter_y],
                alpha = 0.5, color = 'lightgreen', s = 20, edgecolor = 'black'
            )
            axes[1].set_xlabel(scatter_x, fontsize = 7)
            axes[1].set_ylabel(scatter_y, fontsize = 7)
            axes[1].set_title(f"{scatter_x} vs {scatter_y}", fontsize = 8)
        else:
            axes[1].axis('off')
            axes[1].set_title("Scatter Plot Not Available", fontsize = 8)
    else:
        axes[1].axis('off')
        axes[1].set_title("Scatter Plot Not Available", fontsize = 8)

    # Chart 4: Missing values
    nulls = df.isnull().sum()
    nulls = nulls[nulls > 0]
    if not nulls.empty:
        nulls.plot(kind = 'bar', ax = axes[2], color = 'lightcoral', edgecolor = 'black')
        axes[2].set_title("Missing Values by Column", fontsize = 8)
        axes[2].set_ylabel("Count", fontsize = 7)
        axes[2].tick_params(axis = 'x', rotation = 45, labelsize = 7)
    else:
        axes[2].axis('off')
        axes[2].set_title("No Missing Values", fontsize = 8)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format = 'png', dpi = 150, bbox_inches = 'tight')
    buf.seek(0)
    sheet_charts.append(("Summary Charts", buf.getvalue()))
    plt.close(fig)

    return sheet_charts

# --------------- Function to generate Word report with charts ---------------
def generate_word_report(full_report, sheets_data, llm):
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    matplotlib.use('Agg')

    doc = Document()

    # Title and styling 
    title = doc.add_heading('AI Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Report content
    doc.add_heading('Analysis', level = 1)
    for line in full_report.split('\n'):
        line = line.strip()
        if not line:
            pass # Add a blank line for spacing
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level = 2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level = 1)
        elif line.startswith('='*10):
            doc.add_paragraph('─' * 40)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style = 'List Bullet')
        else:
            doc.add_paragraph(line)

    # Charts
    doc.add_page_break()
    doc.add_heading('Auto-Generated Charts', level=1)

    for sheet_name, df in sheets_data.items():
        if len(sheets_data) > 1:
            doc.add_heading(f'Charts - {sheet_name}', level =  2)

        sheet_charts = generate_chart_bytes(df, llm)

        if not sheet_charts:
            doc.add_paragraph("Charts skipped - no meaningful column names found.")
        else:
            for title, img_bytes in sheet_charts:
                buf = io.BytesIO(img_bytes)
                doc.add_paragraph(title).runs[0].bold = True
                doc.add_picture(buf, width = Inches(6))

    # Save the Word document to a BytesIO buffer and return it for download
    doc_buf = io.BytesIO()
    doc.save(doc_buf)
    doc_buf.seek(0)
    return doc_buf

# --------------- Data Agent Tools ---------------
def create_data_agent(df_dict, llm, model_choice):
    """Create a data agent with tools to query the uploaded dataframes. 
    df_dict is a dictionary of {sheet_name: dataframe}"""
    
    # we need to pass df_dict into the tool scope. We do this by creating tools inside this function so they have access to df_dict via closure
    @tool
    def get_dataframe_info() -> str:
        """Use this tool first to understand the uploaded dataset.
        Returns sheet names, column names, data types, shape and sample rows for each sheets."""

        # This shows in the Streamlit UI when tool is called
        print("[AGENT] Calling get_dataframe_info - reading dataset structure")

        info = []
        for sheet_name, df in df_dict.items():
            info.append(f"\n--- Sheet: {sheet_name} ---")
            info.append(f"Shape: {df.shape[0]:,} rows and {df.shape[1]} columns ")
            info.append(f"Columns: {df.columns.tolist()}")
            info.append(f"Data types:\n{df.dtypes.to_string()}")
            info.append(f"Sample rows:\n{df.head(5).to_string()}")
        return "\n".join(info)
    
    @tool
    def query_dataframe(pandas_code: str) -> str:
        """Use this to query the data using pandas code.
        For single sheet data, the dataframe is available as 'df'.
        For multiple sheets, use the sheet name like 'sheets["Sheet1"]'.
        Examples:
        - df['revenue'].sum()
        - df.groupby('category_name')['profit'].mean()
        - df[df['country'] == 'United States']['revenue'].sum()
        - sheets['Sales']['revenue'].sum()
        Always write safe read-only pandas code."""

        print(f"⚡ [AGENT] Running query: {pandas_code[:80]}")

        try:
            # make both df (first sheet) and sheets (all sheets) available
            first_df = list(df_dict.values())[0]
            result = eval(pandas_code,
                {"df": first_df, "sheets": df_dict, "pd": pd}
            )
            return str(result)
        except Exception as e:
            return f"Error: {e}. Try a different pandas expression."

    @tool
    def generate_insights(topic: str) -> str:
        """Generate business insights and recoomendations about a specific topic.
        Use when user asks for analysis, recommendations, or what data means.
        Examples: 
        - 'insights on revenue by country'
        - 'recommendations for improving profit margin'"""
        print(f"💡 [Agent] Generating insights: {topic}")  
        
        try:
            insight_llm = ChatGoogleGenerativeAI(model = model_choice, google_api_key  = api_key)
            messages = [
                SystemMessage(context = """You are senior business analyst.
                              Give 2-3 specific, actionable insights based on the topic.
                              Reference actual numbers where possible.
                              Be concise - max 150 words total."""),
                HumanMessage(contect = f"Dataset columns: {list(df_dict.values())[0].columns.tolist()}\nTopic: {topic}")
            ]  
            response = insight_llm.invoke(messages)
            if isinstance(response.content, list):
                return " ".join(b['text'] for b in response.content if isinstance(b, dict) and b.get('type') == 'text')
            return response.content
        except Exception as e:
            return f"Error generating insights: {str(e)[:100]}"
        
    @tool
    def get_column_values(column_name: str) -> str:
        """Use this tool to see unique values in a specific column.
        helpful for understanding categorical data before filtering.
        Example: get_column_values('country') shows all unique countries."""

        print(f"[AGENT] Checking values in '{column_name}'")

        try:
            first_df = list(df_dict.values())[0]
            unique_values = first_df[column_name].value_counts().head(20)
            return f"Top values in '{column_name}':\n{unique_values.to_string()}"
        except Exception as e:
            return f"Error: {e}"

    @tool
    def generate_chart(chart_request: str) -> str:
        """Use this tool when the user asks for a chart, graph, plot or visualization.
        Describe what chart to generate in chart_request.
        Examples:
        - 'bar chart of revenue by country'
        - 'histogram of anual income'
        - 'scatter plot of revenue vs profit'
        - 'pie chart of orders by category'
        Returns a confirmation message when chart is generated."""

        print(f"[AGENT] Generating chart: {chart_request}")
        
        import base64

        try:
            first_df = list(df_dict.values())[0]
            # ask Gemini to write the matplotlib code for this chart
            chart_llm = ChatGoogleGenerativeAI(model = model_choice, google_api_key = api_key)
            col_info = f"""
                    Columns: {first_df.columns.tolist()}
                    Data types: {first_df.dtypes.to_string()}
                    Sample: {first_df.head(3).to_string()}
                    """
            code_messages = [
                SystemMessage(content="""You are a matplotlib expert.
                              Write Python code to generate the requested chart using matplotlib.
                              The dataframe is available as 'df'.
                              Rules:
                              - Use matplotlib only (no seaborn, no plotly)
                              - Always set a clear title, xlabel, ylabel
                              - Always create the figure with: plt.figure(figsize=(7,5))
                              - Use plt.tight_layout()
                              - Do not call plt.show()
                              - Do not save the file
                              - ALWAYS format large numbers on axes — never use scientific notation
                              - For money/revenue/profit columns use: ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
                              - For count/quantity columns use: ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'{x:,.0f}'))
                              - Always rotate x-axis labels 45 degrees if there are more than 3 categories: plt.xticks(rotation = 45, ha = 'right')
                              - End with: buf = io.BytesIO() 
                              plt.savefig(buf, format = 'png', dpi = 150, bbox_inches = 'tight')
                              buf.seek(0)
                              chart_bytes = buf.getvalue() 
                              plt.close()
                              Respond with only python code, no explaination, no markdown."""),
                HumanMessage(content=f"Dataset info:\n{col_info}\n\nGenerate chart: {chart_request}")
            ]
            code_response = chart_llm.invoke(code_messages)

            # Extract the code
            if isinstance(code_response.content, list):
                chart_code = " ".join(
                    block['text'] for block in code_response.content
                    if isinstance(block, dict) and block.get('type') == 'text'
                )
            else:
                chart_code = code_response.content

            # clean code - remove markdown if present
            chart_code = chart_code.replace('```python', '').replace('```', '').strip()

            # Execute the chart code
            local_vars = {"df": first_df, "pd": pd, "plt": plt, "io": io, "matplotlib": matplotlib}
            exec(chart_code, local_vars)

            # Get the chart bytes
            chart_bytes = local_vars.get('chart_bytes')

            if chart_bytes:
                encoded = base64.b64encode(chart_bytes).decode('utf-8')
                return f"CHART_GENERATED|{encoded}"
            else: 
                return "Chart generation failed - could not extract chart bytes."
            
        except Exception as e:
            return f"Error generating chart: {str(e)}"

    tools = [get_dataframe_info, query_dataframe, get_column_values, generate_insights, generate_chart]
    # Create memory saver
    memory = MemorySaver()
    agent = create_agent(model = llm, tools = tools, checkpointer = memory)
    return agent

# --------------- RAG: Initialize ChromaDB (in-memory) ---------------
@st.cache_resource
def get_rag_resources():
    embeddings = GoogleGenerativeAIEmbeddings(model = "gemini-embedding-001", google_api_key = api_key)

    chroma_client = chromadb.Client()

    splitter = RecursiveCharacterTextSplitter(
        chunk_size = 1000,
        chunk_overlap = 50,
        separators = ["\n\n", "\n", ".", " "]
    )
    return embeddings, chroma_client, splitter

# --------------- RAG: Index a document ---------------
def rag_index_document(uploaded_file, collection, embeddings, splitter):
    file_name = uploaded_file.name
    if file_name in st.session_state.get("rag_indexed_files", []):
        return False, "Already indexed"
    try:
        suffix = os.path.splitext(file_name)[1]
        with tempfile.NamedTemporaryFile(delete = False, suffix = suffix) as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = tmp.name
        
        if suffix == ".pdf":
            loader = PyPDFLoader(tmp_path)
        elif suffix == ".docx":
            loader = Docx2txtLoader(tmp_path)
        elif suffix == ".txt":
            loader = TextLoader(tmp_path)
        else:
            os.unlink(tmp_path)
            return False, f"Unsupported filetype: {suffix}"

        documents = loader.load()
        os.unlink(tmp_path)
        chunks = splitter.split_documents(documents)

        if not chunks:
            return False, "No text could be extracted"

        for i, chunk in enumerate(chunks):
             embedding = embeddings.embed_query(chunk.page_content)
             collection.add(
                 ids = [f"{file_name}_{i}"],
                 embeddings = [embedding],
                 documents = [chunk.page_content],
                 metadatas = [{
                     "source": file_name,
                     "chunk": str(i),
                     "total_chunks": str(len(chunks)),
                     "page": str(chunk.metadata.get("page", ""))
                 }]
             )
        if "rag_indexed_files" not in st.session_state:
            st.session_state.rag_indexed_files = []
        st.session_state.rag_indexed_files.append(file_name)
        return True, len(chunks)
    
    except Exception as e:
        return False, str(e)

# --------------- RAG: Query documents ---------------
def rag_query(question, collection, embeddings, llm, top_k = 4):
    if collection.count() == 0:
        return "No documents indexed yet.", []
    try:
        question_embedding = embeddings.embed_query(question)
        n_results = min(top_k, collection.count())
        results = collection.query(
            query_embeddings = [question_embedding],
            n_results = n_results
        )
        if not results['documents'][0]:
            return "No relevant content found.", []
        
        chunks = results['documents'][0]
        sources = list(set(m['source'] for m in results['metadatas'][0]))
        context = "\n\n---\n\n".join(chunks)

        messages = [
            SystemMessage(content = """You are a helpful assistant that answwers question basede on the provided document excerpts.
                          RULES:
                          - Only use information from the provided content.
                          - If not found say "I could not find this in the uploaded documents."
                          - Always mention which document the answer came from.
                          - Be specific with exact figures when available
                          - Use bullet points when listing multiple items"""),
            HumanMessage(content = f"""Document excerpts:
                         {context}

                         Question: {question}
                        
                         Answer based only on the excerpts above:""")
        ]
        response = llm.invoke(messages)
        if isinstance(response.content, list):
            answer = " ".join(b['text'] for b in response.content
                            if isinstance(b, dict) and b.get('type') == 'text')
        else:
            answer = response.content

        return answer, sources

    except Exception as e:
        return f"Error: {str(e)}", []
    
#  --------------- Session State Initialization ---------------
# Must be initialized before any widget that depends on them
if "report_generated" not in st.session_state: # We use this to keep track of whether a report has been generated yet, so that we can conditionally show the follow-up question section and charts only after a report is generated.
    st.session_state.report_generated = False
if "full_report" not in st.session_state:
    st.session_state.full_report = "" # We store the full report in session state so that we can refer back to it when answering follow-up questions, and also so that we can include it in the Word document download.
if "report_context" not in st.session_state:
    st.session_state.report_context = "" # We use report_context to store the context that we will provide to the model when answering follow-up questions. This can include the original report, and we could also choose to include additional context such as the original data summary or even the raw data if needed. For now, we will just include the full report as the context for follow-up questions, but this can be adjusted based on your needs and experimentation with what provides the best results for follow-up questions.
if "chat_history" not in st.session_state:
    st.session_state.chat_history = [] # We use chat_history to keep track of the conversation history for follow-up questions, so that we can provide that history as context to the model when answering follow-up questions. This allows the model to have the full context of the conversation and provide more coherent and relevant answers to follow-up questions.
if "sheets_data_cache" not in st.session_state:
    st.session_state.sheets_data_cache = {} # We use sheets_data_cache to store the original data from the uploaded sheets in session state, so that we can refer back to the original data when answering follow-up questions and generating charts for follow-up questions. This is important because the user might ask follow-up questions that require referencing the original data, and we want to have that data readily available in session state without needing to re-upload or re-process the file.
if "current_file" not in st.session_state:
    st.session_state.current_file = None # We keep track of the name of the currently uploaded file in session state, so that we can detect when the user uploads a new file and reset the session state accordingly. This is important because if the user uploads a new file, we want to make sure that we clear out any previous report, data, and chat history that was related to the old file, to avoid confusion and ensure that the app is always showing information relevant to the currently uploaded file.
if "charts_generated" not in st.session_state: 
    st.session_state.charts_generated = False # We use this to keep track of whether charts have been generated for the current report, so that we can conditionally show the charts section only after charts have been generated. This is important because generating charts can take additional time, and we want to provide feedback to the user about the status of chart generation, and only show the charts section once the charts are ready.
if "chart_images" not in st.session_state:
    st.session_state.chart_images = {} # We use chart_images to store the generated chart images in session state, so that we can display them in the Streamlit app and also include them in the Word document download. This allows us to keep the generated charts readily available in session state without needing to regenerate them every time we want to display them or include them in the download, which can improve performance and provide a better user experience.
if "agent_charts" not in st.session_state:
    st.session_state.agent_charts = [] # We use agent_charts to store charts that are generated by the data agent in response to user queries. This allows us to display those charts in the Streamlit app and also include them in the Word document download, providing a richer and more interactive experience for the user when they ask follow-up questions that involve chart generation. By keeping these agent-generated charts in session state, we can easily manage and display them as part of the ongoing conversation with the user.
if "agent_thread_id" not in st.session_state:
    st.session_state.agent_thread_id = str(uuid.uuid4()) # We generate a unique thread ID for the data agent's conversation, so that we can keep the conversation history organized and separate from any other interactions in the app. This is important because the data agent might have its own conversation history that is relevant to the user's queries and interactions with the data, and by keeping it organized under a unique thread ID, we can ensure that we are providing the correct context to the model when answering follow-up questions and generating charts based on user queries.
if "data_agent" not in st.session_state:
    st.session_state.data_agent = None # We will store the data agent instance in session state so that we can refer back to it when the user asks follow-up questions that require querying the data or generating charts. This allows us to maintain the state of the data agent and its conversation history across multiple interactions with the user, providing a more seamless and interactive experience when working with the uploaded data and asking questions about it.
if "rag_indexed_files" not in st.session_state:
    st.session_state.rag_indexed_files = [] # We use this to keep track of which files have been indexed in the RAG system, so that we can avoid re-indexing the same file multiple times if the user uploads it again. This is important because indexing can take time and resources, and we want to optimize the performance of the app by only indexing new files that haven't been indexed before. By keeping track of indexed files in session state, we can quickly check if a file has already been indexed and skip the indexing process if it has, providing a more efficient experience for the user when working with the RAG functionality.
if "rag_chat_history" not in st.session_state:
    st.session_state.rag_chat_history = [] # We use rag_chat_history to keep track of the conversation history for the RAG interactions, so that we can provide that history as context to the model when answering questions based on the indexed documents. This allows the model to have the full context of the RAG conversation and provide more coherent and relevant answers to questions that are based on the uploaded documents, enhancing the usefulness of the RAG functionality in the app.
if "rag_session_id" not in st.session_state:
    st.session_state.rag_session_id = str(uuid.uuid4())[:8] # We generate a unique session ID for the RAG interactions, so that we can keep the conversation history organized and separate from any other interactions in the app. This is important because the RAG interactions might have their own conversation history that is relevant to the user's queries about the uploaded documents, and by keeping it organized under a unique session ID, we can ensure that we are providing the correct context to the model when answering questions based on the indexed documents, enhancing the user experience when working with the RAG functionality.

# --------------- Sidebar ---------------
with st.sidebar:
    # Mode selector at top
    st.markdown("### 🔀 Mode")
    app_mode = st.radio(
        "Select mode",
        ["📈 Data Analysis", "🔍 Doc Explore"],
        label_visibility="collapsed"
    )
    st.divider()
 
    if app_mode == "📈 Data Analysis":
        st.header("⚙️ Settings")
        focus_options = {
            "General Insights": "Provide a comprehensive summary of key findings, trends, and recommendations.",
            "Sales & Revenue": "Focus on revenue trends, profit margins, top products, and sales performance.",
            "Customer Analysis": "Focus on customer segments, demographics, behavior, and retention.",
            "HR & Workforce": "Focus on employee demographics, attrition, salary distribution, and performance.",
            "Financial Analysis": "Focus on financial metrics, budget variance, expenses, and forecasting.",
            "Operations": "Focus on operational efficiency, bottlenecks, and process improvements.",
        }
        selected_focus = st.selectbox("What should the report focus on?", list(focus_options.keys()))
        focus = focus_options[selected_focus]
        st.divider()
        st.markdown("**AI Model**")
        model_choice = st.selectbox(
            "Select model",
            ["gemini-3.1-flash-lite", "gemini-3.5-flash"],
            help = "Flash Lite is faster, Flash is more powerful"
        )
        st.caption("ℹ️ Model applies on next report generation")
        st.divider()
        st.markdown("**About**")
        st.markdown("""
        - 🆓 Powered by Google Gemini (free tier)
        - 📋 Supports CSV and Excel files
        - 💬 Ask follow-up questions after report
        - ⬇️ Download as Word or Text
        """)
 
    else:
        # Document Q&A sidebar
        st.header("📁 Upload Documents")
        st.caption("PDF, Word (.docx), and Text files supported")
 
        rag_embeddings, rag_chroma_client, rag_splitter = get_rag_resources()
 
        # Also need model_choice for rag llm
        model_choice = st.selectbox(
            "Select model",
            ["gemini-3.1-flash-lite", "gemini-3.5-flash"],
            help = "Flash Lite is faster, Flash is more powerful"
        )
 
        rag_collection_name = f"rag_{st.session_state.rag_session_id}"
        try:
            rag_collection = rag_chroma_client.get_collection(rag_collection_name)
        except:
            rag_collection = rag_chroma_client.create_collection(rag_collection_name)
 
        rag_files = st.file_uploader(
            "Choose files",
            type = ["pdf", "docx", "txt"],
            accept_multiple_files = True,
            label_visibility = "collapsed"
        )
 
        if rag_files:
            for file in rag_files:
                if file.name not in st.session_state.rag_indexed_files:
                    with st.spinner(f"Indexing {file.name}..."):
                        success, result = rag_index_document(
                            file, rag_collection, rag_embeddings, rag_splitter
                        )
                    if success:
                        st.success(f"✅ {file.name} - {result} chunks")
                    elif result != "Already indexed":
                        st.error(f"❌ {file.name}: {result}")
 
        st.divider()
 
        if st.session_state.rag_indexed_files:
            st.markdown(f"**{len(st.session_state.rag_indexed_files)} document(s) indexed:**")
            for fname in st.session_state.rag_indexed_files:
                st.markdown(f"📄 {fname}")
            st.divider()
            rag_top_k = st.slider("Chunks to retrieve", 1, 8, 4)
            if st.button("🗑️ Clear all documents", width = "stretch"):
                rag_chroma_client.delete_collection(rag_collection_name)
                st.session_state.rag_indexed_files = []
                st.session_state.rag_chat_history = []
                st.session_state.rag_session_id = str(uuid.uuid4())[:8]
                st.rerun()
        else:
            rag_top_k = 4
            st.info("Upload documents to get started")


# --------------- Report Generator Mode ---------------
if app_mode == "📈 Data Analysis":   
    st.title("✦ Clarix - AI Data Analyst")
    st.caption("Generate insights and business recommnedations from your dataset using AI")
    # --------------- Instructions ---------------
    st.markdown("""
    **How it works:**
    1. Upload any CSV or Excel file
    2. Clarix analyzes your data and generates a structured report
    3. Auto-generated charts highlight key trends
    4. Ask follow-up questions about your data
    5. Download the report as a Word document or text file
    """)

    # --------------- File upload ---------------
    uploaded_file = st.file_uploader("Upload your CSV file or Excel file", type=["csv", "xlsx", "xls"])

# Reset session state if no file is uploaded, to clear out any previous report and data when the user removes the uploaded file or uploads a new file. 
# This ensures that the app is always in a clean state when starting to analyze a new dataset, and prevents any confusion from leftover data or reports from previous uploads. 
# It also helps to manage memory usage by clearing out old data that is no longer needed when a new file is uploaded.
    if uploaded_file is None:
        # No file uploaded — reset everything
        st.session_state.report_generated = False
        st.session_state.full_report = ""
        st.session_state.report_context = ""
        st.session_state.chat_history = []
        st.session_state.sheets_data_cache = {} 
        st.session_state.current_file = None
        st.session_state.charts_generated = False
        st.session_state.chart_images = {}
        st.session_state.agent_charts = []
        st.session_state.agent_thread_id = str(uuid.uuid4())
        st.session_state.data_agent = None

# We also keep track of the name of the currently uploaded file in session state, so that we can detect when the user uploads a new file and reset the session state accordingly. 
# This is important because if the user uploads a new file, we want to make sure that we clear out any previous report, data, and chat history that was related to the old file, 
# to avoid confusion and ensure that the app is always showing information relevant to the currently uploaded file.

    if uploaded_file is not None:
    # If a different file is uploaded, reset everything
        if st.session_state.current_file != uploaded_file.name:
            st.session_state.current_file = uploaded_file.name
            st.session_state.report_generated = False
            st.session_state.full_report = ""
            st.session_state.report_context = ""
            st.session_state.chat_history = []
            st.session_state.sheets_data_cache = {} 
            st.session_state.agent_charts = []
            st.session_state.charts_generated = False
            st.session_state.chart_images = {}
            st.session_state.data_agent = None
            st.session_state.agent_thread_id = str(uuid.uuid4())


    if uploaded_file:
        # ----- Load the file
        file_name = uploaded_file.name
        if file_name.endswith(".csv"): 
            # CSV has only 1 sheet. Load it directly.
            df = pd.read_csv(uploaded_file)

            # Check if columns are unnamed (no headers) and fix by using first row as header if needed. 
            # This is important because if there are no column headers, the model will not be able to understand the data and generate meaningful insights, 
            # so we want to make sure that we have proper column headers before proceeding with the analysis. We check if all columns are either unnamed or integers, 
            # which is a common pattern when there are no headers in a CSV file, and if so, we take the first row of the data and set it as the column headers, 
            # and then drop that first row from the data so that it doesn't interfere with the analysis. 
            # This way, we can ensure that we have meaningful column names for the model to reference when generating the report.

            if all(str(col).startswith('Unnamed') or isinstance(col, int) for col in df.columns):
                st.warning("⚠️ No column headers detected. Using first row as headers.")
                df.columns = df.iloc[0]
                df = df.drop(index=0).reset_index(drop=True)
            
            sheets_data = {"Sheet 1": clean_dataframe(df)}

        else:
            x1 = pd.ExcelFile(uploaded_file)
            sheets_names = x1.sheet_names

            if len(sheets_names) == 1:
                # Only 1 sheet, load it directly
                df = x1.parse(sheets_names[0])

                if all(str(col).startswith('Unnamed') or isinstance(col, int) for col in df.columns):
                    st.warning(f"⚠️ Sheet '{sheets_names[0]}' has no column headers. Using first row as headers.")
                    df.columns = df.iloc[0]
                    df = df.drop(index=0).reset_index(drop=True)
                
                sheets_data = {sheets_names[0]: clean_dataframe(df)}

            else:  
                # for multiple sheets, let user select which one to analyze
                st.info(f"This file contains {len(sheets_names)} sheets.")
                selected_sheet = st.multiselect("Select the sheets to analyze:", sheets_names, default = sheets_names[0])

                if not selected_sheet:
                    st.warning("Please select at least one sheet to proceed.")
                    st.stop()
                
                sheets_data = {}
                for sheet in selected_sheet:
                    df = x1.parse(sheet)
                    if all(str(col).startswith('Unnamed') or isinstance(col, int) for col in df.columns):
                        st.warning(f"⚠️ Sheet '{sheet}' has no column headers. Using first row as headers.")
                        df.columns = df.iloc[0]
                        df = df.drop(index=0).reset_index(drop=True)
                    sheets_data[sheet] = clean_dataframe(df)

                for sheet_name, df in sheets_data.items():
                    row_count = df.shape[0]
                    if row_count > 500000:
                        st.warning(f"Sheet '{sheet_name}' contains {row_count:,} rows, which may impact performance.")
                    elif row_count > 100000:
                        st.info(f"Sheet '{sheet_name}' contains {row_count:,} rows. Performance may be slower with larger datasets.")

        # ----- Show quick stats and Preview the data for each sheet
        for sheet_name, df in sheets_data.items():
            if len(sheets_data) > 1:
                st.subheader(f"Sheet: {sheet_name}")

            col1, col2, col3, col4, col5 = st.columns(5)
            col1.metric("Rows", f"{df.shape[0]:,}")
            col2.metric("Columns", f"{df.shape[1]:,}")
            col3.metric("Numeric Columns", f"{df.select_dtypes(include='number').shape[1]}")
            col4.metric("Categorical Columns", f"{df.select_dtypes(include=['object', 'string']).shape[1]}")
            col5.metric("Missing Values", f"{df.isnull().sum().sum():,}")

            with st.expander(f"Preview - {sheet_name}", expanded = True):
                st.dataframe(df.head(10), width = "stretch")

        st.divider()

        # ----- Generate Button ------

        # When the user clicks the button, we will generate the report using Gemini and display it on the page. 
        # We will also generate some auto-charts based on the data and include them in the report.
        # We will use session state to keep track of the generated report and the conversation history for follow-up questions.
        if st.button("Generate Report", type = "primary", width = "stretch"):
            # Reset session state for new report generation
        # Reset for new report
            st.session_state.report_generated = False 
            st.session_state.charts_generated = False
            st.session_state.chat_history = []
            st.session_state.full_report = ""
            st.session_state.report_context = ""
            st.session_state.sheets_data_cache = sheets_data

            # --- Generate the report using Gemini
            with st.spinner("Analyzing your data with Gemini... (~15-30 seconds)"):
                # Initialize the llm
                llm = ChatGoogleGenerativeAI(model = model_choice, google_api_key = api_key)
                # Generate a report for each sheet and combine them into one final report
                sheet_reports = {}

                # Loop through each sheet and generate a report for it
                for sheet_name, df in sheets_data.items():
                    # Summarize the data for the current sheet
                    summary = summarize_dataframe(df) 

                    # Create the prompt
                    messages = [
                        SystemMessage(content = f"""You are a senior data analyst. You are given a summary of a dataset. 
                                    Your task is to analyze the summary and provide insights, trends, and potential business recommendations based on the data.
                                    Return the report with these sections:
                                    ## Executive Summary
                                    ## Key Findings
                                    ## Data Quality Notes
                                    ## Recommendations
                                    Focus on: {focus}"""),
                    
                        HumanMessage(content = f"Analyze this data from '{sheet_name}':\n\n{summary}")
                    ]
                    try:
                        response = llm.invoke(messages) 
                        sheet_reports[sheet_name] = response.content[0]['text']
                    except Exception as e:
                        st.error(f"Error occurred while generating report for sheet '{sheet_name}': {e}. Please try again.")
                        st.stop()

                # Combine all sheet reports into one final report
                if len(sheet_reports) == 1:
                    # If only one sheet, no need to add sheet names in the report
                    full_report = list(sheet_reports.values())[0]
                else:
                    # If multiple sheets, add sheet names as headers in the report
                    full_report = ""
                    for sheet_name, report in sheet_reports.items():
                        full_report += f"\n\n{'='*50}\n"
                        full_report += f"\nReport for Sheet: {sheet_name}\n"
                        full_report += f"\n{'='*50}\n\n"
                        full_report += report

                # Store the full report and sheets data in session state so that we can refer back to it when answering follow-up questions and generating charts for follow-up questions
                st.session_state.full_report = full_report
                st.session_state.report_context = full_report  # We use the full report as the context for follow-up questions, but this can be adjusted based on your needs and experimentation with what provides the best results for follow-up questions. For example, you could choose to include the original data summary or even the raw data as part of the context for follow-up questions if you find that leads to better answers from the model.
                st.session_state.report_generated = True  # We set this to True so that we can conditionally show the follow-up question section and charts only after a report has been generated. This helps to keep the UI clean and focused, and only show the relevant sections when they are applicable.
            st.rerun()
        
        # ----- Display report
        if st.session_state.report_generated: # We only show the report and charts if a report has been generated, to keep the UI clean and focused.

            llm = ChatGoogleGenerativeAI(model = model_choice, google_api_key = api_key)

            st.success("Report generated successfully!")
            st.markdown("---")
            st.markdown(st.session_state.full_report)

            # --- Only generate charts if not already generated 
            st.markdown("---")
            if not st.session_state.charts_generated:
                chart_images = {}
                
                for sheet_name, df in st.session_state.sheets_data_cache.items():
                    sheet_charts = generate_chart_bytes(df, llm)
                    chart_images[sheet_name] = sheet_charts

                st.session_state.chart_images = chart_images
                st.session_state.charts_generated = True
                st.rerun()
            
            if st.session_state.charts_generated:
                for sheet_name, charts in st.session_state.chart_images.items():
                    if len(st.session_state.sheets_data_cache) > 1:
                        st.subheader(f"Auto Generated Charts — {sheet_name}")
                    else:
                        st.subheader("Auto Generated Charts")

                    if not charts:
                        st.info("⚠️ Charts skipped — no meaningful column names found.")
                    else:
                        for title, img_bytes in charts:
                            st.markdown(f"**{title}**")
                            st.image(img_bytes, width = "stretch")

            # Download buttons for both Word and Text formats
            st.markdown("---")
            st.markdown("### Download Report")
            base_name = os.path.splitext(uploaded_file.name)[0]

            # Provide both Word and Text download options side by side
            col1, col2 = st.columns([1, 1])
            # Text file download
            with col1:            
                st.download_button(
                label = "Download Report as Text File",
                data = st.session_state.full_report,
                file_name = f"{base_name}_Report.txt",
                mime = "text/plain",
                width = "stretch"
                )
            # Word document download
            with col2:
                doc_buf = generate_word_report(st.session_state.full_report, st.session_state.sheets_data_cache, llm)
                st.download_button(
                label = "Download Report as Word Document",
                data = doc_buf,
                file_name = f"{base_name}_Report.docx",
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                width = "stretch"
                )

            # --- Follow-up Questions
            # --- Data Agent Chat
            st.markdown("---")
            st.markdown("### Ask anything about your data")

            # Display existing chat history
            for message in st.session_state.chat_history:
                if message["role"] == "user":
                    with st.chat_message("user"):
                        st.markdown(message['content'])
                else:
                    with st.chat_message("assistant"):
                        st.markdown(message['content'])
                        if 'chart' in message:
                            st.image(message['chart'], width = 'stretch')

            # Chat input
            user_input = st.chat_input("Ask anything about your data...")
            if user_input:
                chart_bytes = None
                with st.chat_message("user"):
                    st.markdown(user_input)

                with st.chat_message("assistant"):
                    result = None
                    ai_response = ""

                    with st.status("🤔 Thinking...", expanded = True) as status:
                        try:
                            if st.session_state.data_agent is None:
                                status.update(label = "⚙️ Setting up agent...")
                                st.session_state.data_agent = create_data_agent(st.session_state.sheets_data_cache, llm, model_choice)
                        
                            agent = st.session_state.data_agent

                            # Build message history for context
                            agent_messages = [HumanMessage(content = user_input)]

                            # Add recent chat history for context
                            recent_history = st.session_state.chat_history[-6:]
                            history_context = ""
                            if recent_history:
                                history_context = "\n\nPrevious conversation:\n"
                                for msg in recent_history:
                                    role = "User" if msg["role"] == "user" else "Assistant"
                                    history_context += f"{role}: {msg['content']}\n"

                            # Include report context and chat history in the question
                            full_question = f"""You are a data analyst assistant. 
                            Here is the report that was generated earlier for context:
                            {st.session_state.report_context[:2000]}

                            Previous conversation:
                            {history_context}

                            Now answer this question by querying the raw data:
                            {user_input}

                            IMPORTANT RULES:
                            - Only call generate_chart tool if the user EXPLICITLY asks for a chart, graph, plot or visualization
                            - If the user just asks a data question, answer with numbers and text only
                            - Do NOT generate charts automatically
                            - Use query_dataframe tool to get exact numbers
                            - Be concise and specific"""

                            result = agent.invoke({"messages": [HumanMessage(content = full_question)]},
                                            config = {"configurable": {"thread_id": st.session_state.agent_thread_id}})
                        
                            # Create callback with status container
                            callback = StreamlitAgentCallback(status)

                            status.update(label = "🤔 Thinking...")

                            result = agent.invoke({"messages": [HumanMessage(content=full_question)]},
                            config={
                                "configurable": {"thread_id": st.session_state.agent_thread_id},
                                "callbacks": [callback]  # passing callback here
                            })

                            # Handle both string and list content formats
                            last_message = result["messages"][-1]
                            content = last_message.content

                            if isinstance(content, list):
                                # Extract text from list format
                                ai_response = " ".join(
                                    block['text'] for block in content 
                                    if isinstance(block, dict) and block.get('type') == 'text'
                                )
                            elif isinstance(content, str):
                                ai_response = content
                            else:
                                ai_response = str(content)

                            if result:
                                for msg in result["messages"]:
                                    if hasattr(msg, 'content') and isinstance(msg.content, str):
                                        if msg.content.startswith("CHART_GENERATED|"):
                                            import base64
                                            encoded = msg.content.split("|", 1)[1]
                                            chart_bytes = base64.b64decode(encoded)
                                            break
                            
                                status.update(label = "Done!", state = "complete", expanded = False)
                        except Exception as e:
                            ai_response = f"Sorry, I encountered an error: {str(e)}"
                            status.update(label = "Error", state = "error", expanded = True)
                            st.error(ai_response)

                    st.markdown(ai_response)

                # Dispaly any charts generated by the agent
                    if chart_bytes:
                        st.image(chart_bytes, width = 600)

                # Save to chat history
                st.session_state.chat_history.append({"role": "user", "content": user_input})
                if chart_bytes:
                    st.session_state.chat_history.append({"role": "assistant", "content": ai_response, "chart": chart_bytes})
                else:
                    st.session_state.chat_history.append({"role": "assistant", "content": ai_response})

# --------------- Document Q&A Mode ---------------
else:
    st.title("📄 Doc Explore")
    st.caption("Search and extract insights from your documents instantly")
 
    if not st.session_state.rag_indexed_files:
        st.info("👈 Upload documents in the sidebar to get started")
        st.markdown("""
**What you can do:**
- 📄 Upload multiple PDFs, Word docs, or text files
- 💬 Ask questions across all documents at once
- 🔍 Get answers with source citations
- 🧠 Search by meaning - not just keywords
        """)
    else:
        rag_llm = ChatGoogleGenerativeAI(model = "gemini-2.5-flash-lite", google_api_key = api_key)
 
        st.success(f"**{len(st.session_state.rag_indexed_files)} document(s) ready** - ask anything!")
 
        # Chat history
        for message in st.session_state.rag_chat_history:
            if message["role"] == "user":
                with st.chat_message("user"):
                    st.markdown(message["content"])
            else:
                with st.chat_message("assistant"):
                    st.markdown(message["content"])
                    if message.get("sources"):
                        st.caption(f"📄 Sources: {', '.join(message['sources'])}")
 
        # Chat input
        rag_question = st.chat_input("Ask a question about your documents...")
        if rag_question:
            with st.chat_message("user"):
                st.markdown(rag_question)
 
            with st.chat_message("assistant"):
                with st.status("🔍 Searching documents...", expanded = False) as status:
                    answer, sources = rag_query(
                        rag_question, rag_collection, rag_embeddings, rag_llm, rag_top_k
                    )
                    status.update(label="✅ Done!", state = "complete")
 
                st.markdown(answer)
                if sources:
                    st.caption(f"📄 Sources: {', '.join(sources)}")
 
            st.session_state.rag_chat_history.append({"role": "user", "content": rag_question})
            st.session_state.rag_chat_history.append({"role": "assistant", "content": answer, "sources": sources})